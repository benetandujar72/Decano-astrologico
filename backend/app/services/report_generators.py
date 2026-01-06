"""
Generadores de informes astrológicos en múltiples formatos
Soporta: PDF, DOCX, Markdown, HTML
"""
from typing import Dict, Optional
from io import BytesIO
import json
import re
from datetime import datetime
import base64
import os

# Generador de imágenes de carta astral
try:
    from app.services.chart_image_generator import generate_chart_image
    CHART_IMAGE_AVAILABLE = True
except ImportError:
    CHART_IMAGE_AVAILABLE = False
    print("⚠️ Generador de imágenes no disponible")

# Generador de portadas
try:
    from app.services.report_cover_generator import generate_mystical_cover, add_cover_to_pdf
    COVER_AVAILABLE = True
except ImportError:
    COVER_AVAILABLE = False
    print("⚠️ Generador de portadas no disponible")

# PDF Generation
try:
    from reportlab.lib.pagesizes import letter, A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.platypus import (
        SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle,
        PageBreak, Image as RLImage, KeepTogether
    )
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_JUSTIFY
    from reportlab.pdfgen import canvas
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False
    print("⚠️ ReportLab no disponible. Instala con: pip install reportlab")

# DOCX Generation
try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("⚠️ python-docx no disponible. Instala con: pip install python-docx")


class ReportGenerator:
    """Generador principal de informes astrológicos"""
    
    def __init__(self, carta_data: Dict, analysis_text: Optional[str] = None, nombre: str = ""):
        """
        Args:
            carta_data: Datos completos de la carta astral
            analysis_text: Texto del análisis psico-astrológico (opcional)
            nombre: Nombre del consultante (para portada)
        """
        self.carta = carta_data
        self.analysis = analysis_text or "Análisis no disponible"
        self.datos = carta_data.get('datos_entrada', {})
        self.planetas = carta_data.get('planetas', {})
        self.casas = carta_data.get('casas', [])
        self.angulos = carta_data.get('angulos', {})
        self.nombre = nombre

        # Normalizar campos entre versiones (compatibilidad)
        self.fecha_local = self.datos.get('fecha_local') or self.datos.get('fecha') or 'N/A'
        self.hora_local = self.datos.get('hora_local') or self.datos.get('hora') or 'N/A'
        self.zona_horaria = self.datos.get('zona_horaria') or 'UTC'
        self.fecha_utc = self.datos.get('fecha_utc') or 'N/A'

        # Inyectar nombre al payload para que otros generadores puedan usarlo si quieren
        try:
            if self.nombre and not self.datos.get('nombre'):
                self.datos['nombre'] = self.nombre
        except Exception:
            pass
        
        # Generar imagen de carta astral
        self.chart_image = None
        if CHART_IMAGE_AVAILABLE:
            try:
                self.chart_image = generate_chart_image(carta_data, size=(800, 800))
            except Exception as e:
                print(f"⚠️ Error generando imagen de carta: {e}")
        
        # Generar portada mística
        self.cover_image = None
        if COVER_AVAILABLE and nombre:
            try:
                sol_signo = self.planetas.get('Sol', {}).get('signo', '')
                luna_signo = self.planetas.get('Luna', {}).get('signo', '')
                asc_signo = self.angulos.get('ascendente', {}).get('signo', '')
                
                self.cover_image = generate_mystical_cover(
                    nombre=nombre,
                    fecha=self.fecha_local if self.fecha_local != 'N/A' else '',
                    hora=self.hora_local if self.hora_local != 'N/A' else '',
                    lugar=f"Lat {self.datos.get('latitud', 0)}, Lon {self.datos.get('longitud', 0)}",
                    tipo_analisis="Carta Natal Completa",
                    ascendente=asc_signo,
                    sol_signo=sol_signo,
                    luna_signo=luna_signo
                )
            except Exception as e:
                print(f"⚠️ Error generando portada: {e}")
    
    # ===================== GENERADOR MARKDOWN =====================
    
    def generate_markdown(self) -> str:
        """Genera informe en formato Markdown"""
        md = f"""# 🌟 Carta Astral Completa

## 📋 Datos Personales

- **Fecha:** {self.fecha_local} {self.hora_local}
- **Ubicación:** Lat {self.datos.get('latitud', 0)}, Lon {self.datos.get('longitud', 0)}
- **Zona Horaria:** {self.zona_horaria}
- **Fecha UTC:** {self.fecha_utc}

---

## 🪐 Posiciones Planetarias

| Planeta | Posición | Casa | Retrogradación |
|---------|----------|------|----------------|
"""
        
        for nombre, pos in self.planetas.items():
            if pos:
                retro = "✓ Retrógrado" if pos.get('retrogrado', False) else ""
                casa = pos.get('casa', '?')
                md += f"| {nombre} | {pos['texto']} | Casa {casa} | {retro} |\n"
        
        # Parte de Fortuna
        pf = self.angulos.get('parte_fortuna', {})
        md += f"| Parte de Fortuna | {pf.get('texto', 'N/A')} | - | - |\n"
        
        md += f"""
---

## 🔺 Ángulos Principales

- **Ascendente:** {self.angulos.get('ascendente', {}).get('texto', 'N/A')}
- **Medio Cielo:** {self.angulos.get('medio_cielo', {}).get('texto', 'N/A')}

---

## 🏠 Cúspides de Casas (Placidus)

| Casa | Cúspide |
|------|---------|
"""
        
        for casa in self.casas:
            md += f"| Casa {casa['numero']} | {casa['texto']} |\n"
        
        md += f"""
---

## 📖 Análisis Psico-Astrológico

{self.analysis}

---

*Informe generado el {datetime.now().strftime('%d/%m/%Y %H:%M')}*
"""
        
        return md
    
    # ===================== GENERADOR HTML =====================
    
    def generate_html(self) -> str:
        """Genera informe HTML usando plantillas editables (Jinja2)."""
        return self._render_template_html()

    def _render_template_html(self) -> str:
        """Renderiza HTML usando `backend/app/templates/report/` (template.html + styles.css)."""
        try:
            from jinja2 import Environment, FileSystemLoader, select_autoescape
        except Exception:
            # Fallback: HTML simple (sin plantillas) basado en markdown
            return f"<pre>{self.generate_markdown()}</pre>"

        template_dir = os.getenv("REPORT_TEMPLATE_DIR") or os.path.join(os.path.dirname(__file__), "..", "templates", "report")
        template_dir = os.path.abspath(template_dir)

        env = Environment(
            loader=FileSystemLoader(template_dir),
            autoescape=select_autoescape(["html", "xml"]),
        )
        tpl = env.get_template("template.html")

        # Textos editables (archivo opcional)
        texts = {
            "personal_data_title": "Datos Personales",
            "chart_title": "Carta Astral",
            "planets_title": "Posiciones Planetarias",
            "houses_title": "Cúspides de Casas (Placidus)",
            "analysis_title": "Análisis Psico-Astrológico",
            "brand": "Motor Fraktal",
            "subtitle": "Informe técnico de astropsicología",
            "footer_left": "Informe generado por Motor Fraktal",
            "footer_right": datetime.now().strftime("%d/%m/%Y"),
            "title": "Informe Astrológico",
        }
        try:
            texts_path = os.getenv("REPORT_TEXTS_PATH") or os.path.join(template_dir, "config_texts.json")
            if os.path.exists(texts_path):
                with open(texts_path, "r", encoding="utf-8") as f:
                    texts.update(json.load(f) or {})
        except Exception:
            pass

        planetas = []
        for nombre, pos in (self.planetas or {}).items():
            if not pos:
                continue
            planetas.append(
                {
                    "name": nombre,
                    "texto": pos.get("texto", ""),
                    "casa": pos.get("casa", "?"),
                    "retrogrado": bool(pos.get("retrogrado", False)),
                }
            )

        etiquetas_casas = {"1": "Ascendente (AC)", "4": "Fondo Cielo (IC)", "7": "Descendente (DC)", "10": "Medio Cielo (MC)"}
        casas = []
        for casa in (self.casas or []):
            num = str(casa.get("numero"))
            casas.append({"numero": num, "texto": casa.get("texto", ""), "etiqueta": etiquetas_casas.get(num, "-")})

        pf = self.angulos.get("parte_fortuna", {}) if isinstance(self.angulos, dict) else {}

        # Markdown -> HTML básico (párrafos + headers)
        analysis_html = ""
        try:
            parts = (self.analysis or "").split("\n\n")
            html_parts = []
            for p in parts:
                p = p.strip()
                if not p:
                    continue
                if p.startswith("### "):
                    html_parts.append(f"<h3>{self._format_inline(p[4:])}</h3>")
                elif p.startswith("## "):
                    html_parts.append(f"<h2>{self._format_inline(p[3:])}</h2>")
                elif p.startswith("# "):
                    html_parts.append(f"<h1>{self._format_inline(p[2:])}</h1>")
                else:
                    html_parts.append(f"<p>{self._format_inline(p)}</p>")
            analysis_html = "\n".join(html_parts)
        except Exception:
            analysis_html = f"<p>{(self.analysis or '').replace(chr(10), '<br/>')}</p>"

        chart_data_uri = None
        if self.chart_image:
            try:
                self.chart_image.seek(0)
                img_base64 = base64.b64encode(self.chart_image.read()).decode("utf-8")
                chart_data_uri = f"data:image/png;base64,{img_base64}"
            except Exception:
                chart_data_uri = None

        return tpl.render(
            title=texts.get("title", "Informe Astrológico"),
            brand=texts.get("brand", "Motor Fraktal"),
            subtitle=texts.get("subtitle", ""),
            nombre=self.nombre or (self.datos.get("nombre") if isinstance(self.datos, dict) else "") or "Consultante",
            fecha_local=self.fecha_local,
            hora_local=self.hora_local,
            zona_horaria=self.zona_horaria,
            fecha_utc=self.fecha_utc,
            latitud=self.datos.get("latitud", ""),
            longitud=self.datos.get("longitud", ""),
            generated_at=datetime.now().strftime("%d/%m/%Y %H:%M"),
            footer_left=texts.get("footer_left", ""),
            footer_right=texts.get("footer_right", ""),
            planetas=planetas,
            casas=casas,
            parte_fortuna={"texto": pf.get("texto", "")} if pf else None,
            analysis_html=analysis_html,
            chart_image_data_uri=chart_data_uri,
            texts=texts,
        )
    
    # ===================== GENERADOR PDF =====================
    
    def _format_inline(self, text: str) -> str:
        """Aplica formato inline (negrita, cursiva)"""
        # Negrita: **texto** -> <b>texto</b>
        text = re.sub(r'\*\*(.*?)\*\*', r'<b>\1</b>', text)
        # Cursiva: *texto* -> <i>texto</i> (evitando bullets al inicio)
        text = re.sub(r'(?<!\*)\*(?!\*)(.*?)\*', r'<i>\1</i>', text)
        
        # Saltos de línea
        text = text.replace('\n', '<br/>')
        return text

    def _parse_markdown(self, text: str, styles) -> list:
        """Convierte texto Markdown básico a objetos ReportLab"""
        story_elements = []
        
        # Estilos
        normal_style = ParagraphStyle(
            'CustomNormal',
            parent=styles['Normal'],
            fontSize=11,
            spaceAfter=12,
            alignment=TA_JUSTIFY,
            leading=14
        )
        
        h1_style = ParagraphStyle(
            'CustomH1',
            parent=styles['Heading1'],
            fontSize=18,
            textColor=colors.HexColor('#667eea'),
            spaceAfter=12,
            spaceBefore=12,
            fontName='Helvetica-Bold'
        )
        
        h2_style = ParagraphStyle(
            'CustomH2',
            parent=styles['Heading2'],
            fontSize=16,
            textColor=colors.HexColor('#764ba2'),
            spaceAfter=10,
            spaceBefore=10,
            fontName='Helvetica-Bold'
        )
        
        h3_style = ParagraphStyle(
            'CustomH3',
            parent=styles['Heading3'],
            fontSize=14,
            textColor=colors.HexColor('#4a5568'),
            spaceAfter=8,
            spaceBefore=8,
            fontName='Helvetica-Bold'
        )
        
        bullet_style = ParagraphStyle(
            'CustomBullet',
            parent=normal_style,
            leftIndent=20,
            firstLineIndent=0,
            spaceAfter=6,
            bulletIndent=10
        )

        # Procesar bloques
        blocks = text.split('\n\n')
        
        for block in blocks:
            block = block.strip()
            if not block:
                continue
                
            # Detectar Headers
            if block.startswith('### '):
                content = block[4:]
                content = self._format_inline(content)
                story_elements.append(Paragraph(content, h3_style))
            elif block.startswith('## '):
                content = block[3:]
                content = self._format_inline(content)
                story_elements.append(Paragraph(content, h2_style))
            elif block.startswith('# '):
                content = block[2:]
                content = self._format_inline(content)
                story_elements.append(Paragraph(content, h1_style))
            elif block.startswith('---'):
                story_elements.append(Spacer(1, 0.2*inch))
            else:
                # Detectar listas
                lines = block.split('\n')
                is_list = False
                for line in lines:
                    if line.strip().startswith('* ') or line.strip().startswith('- '):
                        is_list = True
                        break
                
                if is_list:
                    for line in lines:
                        line = line.strip()
                        if line.startswith('* ') or line.startswith('- '):
                            content = line[2:]
                            content = self._format_inline(content)
                            story_elements.append(Paragraph(f"• {content}", bullet_style))
                        else:
                            content = self._format_inline(line)
                            story_elements.append(Paragraph(content, normal_style))
                else:
                    # Párrafo normal
                    content = self._format_inline(block)
                    story_elements.append(Paragraph(content, normal_style))
            
            story_elements.append(Spacer(1, 0.05*inch))
            
        return story_elements

    def generate_pdf(self) -> BytesIO:
        """Genera informe en formato PDF"""
        import sys
        
        if not PDF_AVAILABLE:
            raise ImportError("ReportLab no está instalado. Instala con: pip install reportlab")
        
        try:
            # Motor preferido: HTML + CSS (WeasyPrint) si está disponible
            if (os.getenv("PDF_ENGINE") or "").lower().strip() in {"html", "weasyprint"}:
                try:
                    from weasyprint import HTML, CSS  # type: ignore
                    template_dir = os.getenv("REPORT_TEMPLATE_DIR") or os.path.join(os.path.dirname(__file__), "..", "templates", "report")
                    template_dir = os.path.abspath(template_dir)
                    css_path = os.getenv("REPORT_STYLES_PATH") or os.path.join(template_dir, "styles.css")
                    html_str = self._render_template_html()
                    stylesheets = [CSS(filename=css_path)] if os.path.exists(css_path) else None
                    pdf_bytes = HTML(string=html_str, base_url=template_dir).write_pdf(stylesheets=stylesheets)
                    return BytesIO(pdf_bytes)
                except Exception as e:
                    print(f"⚠️ PDF_ENGINE=html falló, usando ReportLab: {type(e).__name__}: {e}", file=sys.stderr)

            buffer = BytesIO()
            doc = SimpleDocTemplate(buffer, pagesize=A4, 
                                    topMargin=0.75*inch, bottomMargin=0.75*inch)
            
            # Estilos
            styles = getSampleStyleSheet()
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Heading1'],
                fontSize=24,
                textColor=colors.HexColor('#667eea'),
                spaceAfter=30,
                alignment=TA_CENTER,
                fontName='Helvetica-Bold'
            )
            
            heading_style = ParagraphStyle(
                'CustomHeading',
                parent=styles['Heading2'],
                fontSize=16,
                textColor=colors.HexColor('#764ba2'),
                spaceAfter=12,
                spaceBefore=20,
                fontName='Helvetica-Bold'
            )
            
            normal_style = ParagraphStyle(
                'CustomNormal',
                parent=styles['Normal'],
                fontSize=11,
                spaceAfter=12,
                alignment=TA_JUSTIFY
            )
            
            # Contenido
            story = []
            
            # Título
            story.append(Paragraph("🌟 Carta Astral Completa", title_style))
            story.append(Spacer(1, 0.3*inch))
            
            # Datos Personales
            story.append(Paragraph("📋 Datos Personales", heading_style))
            datos_text = f"""
            <b>Fecha:</b> {self.fecha_local} {self.hora_local}<br/>
            <b>Ubicación:</b> Lat {self.datos.get('latitud', 0)}, Lon {self.datos.get('longitud', 0)}<br/>
            <b>Zona Horaria:</b> {self.zona_horaria}<br/>
            <b>Fecha UTC:</b> {self.fecha_utc}
            """
            story.append(Paragraph(datos_text, normal_style))
            story.append(Spacer(1, 0.2*inch))
            
            # Carta Astral Visual
            if self.chart_image:
                try:
                    story.append(Paragraph("🌟 Carta Astral Visual", heading_style))
                    self.chart_image.seek(0)
                    img = RLImage(self.chart_image, width=4.5*inch, height=4.5*inch)
                    story.append(img)
                    story.append(Spacer(1, 0.3*inch))
                except Exception as e:
                    print(f"⚠️ Error añadiendo imagen al PDF: {e}", file=sys.stderr)
            
            # Tabla de Planetas
            story.append(Paragraph("🪐 Posiciones Planetarias (Efemérides)", heading_style))
            
            planetas_data = [['Planeta', 'Longitud Eclíptica', 'Casa', 'Estado']]
            for nombre, pos in self.planetas.items():
                if pos:
                    retro = 'Retrógrado' if pos.get('retrogrado', False) else ''
                    casa = f"Casa {pos.get('casa', '?')}"
                    planetas_data.append([nombre, pos['texto'], casa, retro])
            
            # Añadir Parte de Fortuna
            pf = self.angulos.get('parte_fortuna', {})
            planetas_data.append(['Parte de Fortuna', pf.get('texto', 'N/A'), '-', '-'])
            
            planetas_table = Table(planetas_data, colWidths=[1.5*inch, 2*inch, 1*inch, 1.5*inch])
            planetas_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#667eea')),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 12),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                ('GRID', (0, 0), (-1, -1), 1, colors.grey),
                ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#f8f9fa')]),
            ]))
            
            story.append(planetas_table)
            story.append(Spacer(1, 0.3*inch))
            
            # Ángulos y Cúspides
            story.append(Paragraph("🔺 Ángulos y Cúspides de Casas", heading_style))
            
            # Preparar datos de casas con etiquetas especiales
            casas_data = [['Casa', 'Cúspide', 'Significado']]
            etiquetas_casas = {
                '1': 'Ascendente (AC)',
                '4': 'Fondo Cielo (IC)',
                '7': 'Descendente (DC)',
                '10': 'Medio Cielo (MC)'
            }
            
            for casa in self.casas:
                num = str(casa['numero'])
                etiqueta = etiquetas_casas.get(num, '-')
                casas_data.append([f"Casa {num}", casa['texto'], etiqueta])
            
            casas_table = Table(casas_data, colWidths=[1.5*inch, 3*inch, 2*inch])
            casas_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#667eea')),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 12),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                ('GRID', (0, 0), (-1, -1), 1, colors.grey),
                ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#f8f9fa')]),
            ]))
            
            story.append(casas_table)
            story.append(PageBreak())
            
            # Análisis (USANDO PARSER MARKDOWN)
            story.append(Paragraph("📖 Análisis Psico-Astrológico", heading_style))
            
            # Usar el nuevo parser
            markdown_elements = self._parse_markdown(self.analysis, styles)
            story.extend(markdown_elements)
            
            # Footer
            story.append(Spacer(1, 0.3*inch))
            footer_text = f"<i>Informe generado el {datetime.now().strftime('%d/%m/%Y %H:%M')}</i>"
            footer_style = ParagraphStyle('Footer', parent=styles['Normal'], 
                                          fontSize=9, textColor=colors.grey, alignment=TA_CENTER)
            story.append(Paragraph(footer_text, footer_style))
        
            # Construir PDF
            try:
                doc.build(story)
                buffer.seek(0)
            except Exception as build_err:
                print(f"❌ Error construyendo PDF: {build_err}", file=sys.stderr)
                import traceback
                traceback.print_exc(file=sys.stderr)
                raise Exception(f"Error al construir el PDF: {str(build_err)}")
            
            # Añadir portada mística al inicio
            if self.cover_image and COVER_AVAILABLE:
                try:
                    final_buffer = add_cover_to_pdf(buffer, self.cover_image)
                    return final_buffer
                except Exception as e:
                    print(f"⚠️ Error añadiendo portada al PDF: {e}", file=sys.stderr)
                    buffer.seek(0)
                    return buffer
            
            return buffer
        except Exception as e:
            print(f"❌ Error en generate_pdf: {type(e).__name__}: {e}", file=sys.stderr)
            import traceback
            traceback.print_exc(file=sys.stderr)
            raise
    
    # ===================== GENERADOR DOCX =====================
    
    def generate_docx(self) -> BytesIO:
        """Genera informe en formato DOCX"""
        import sys
        
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx no está instalado. Instala con: pip install python-docx")
        
        try:
            doc = Document()
            
            # Configurar estilos del documento
            style = doc.styles['Normal']
            style.font.name = 'Calibri'
            style.font.size = Pt(11)
            
            # Título
            title = doc.add_heading('🌟 Carta Astral Completa', level=0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title.runs[0].font.color.rgb = RGBColor(102, 126, 234)
            
            doc.add_paragraph()
            
            # Datos Personales
            doc.add_heading('📋 Datos Personales', level=1)
            doc.add_paragraph(f"Fecha: {self.fecha_local} {self.hora_local}")
            doc.add_paragraph(f"Ubicación: Lat {self.datos.get('latitud', 0)}, Lon {self.datos.get('longitud', 0)}")
            doc.add_paragraph(f"Zona Horaria: {self.zona_horaria}")
            doc.add_paragraph(f"Fecha UTC: {self.fecha_utc}")
            
            # Posiciones Planetarias
            doc.add_heading('🪐 Posiciones Planetarias', level=1)
            
            table = doc.add_table(rows=1, cols=4)
            table.style = 'Light Grid Accent 1'
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Planeta'
            hdr_cells[1].text = 'Posición'
            hdr_cells[2].text = 'Casa'
            hdr_cells[3].text = 'Estado'
            
            for nombre, pos in self.planetas.items():
                if pos:
                    row_cells = table.add_row().cells
                    row_cells[0].text = nombre
                    row_cells[1].text = pos['texto']
                    row_cells[2].text = f"Casa {pos.get('casa', '?')}"
                    row_cells[3].text = 'Retrógrado' if pos.get('retrogrado', False) else ''
            
            # Parte de Fortuna
            pf = self.angulos.get('parte_fortuna', {})
            row_cells = table.add_row().cells
            row_cells[0].text = 'Parte de Fortuna'
            row_cells[1].text = pf.get('texto', 'N/A')
            row_cells[2].text = '-'
            row_cells[3].text = '-'
            
            doc.add_paragraph()
            
            # Ángulos
            doc.add_heading('🔺 Ángulos Principales', level=1)
            doc.add_paragraph(f"Ascendente: {self.angulos.get('ascendente', {}).get('texto', 'N/A')}")
            doc.add_paragraph(f"Medio Cielo: {self.angulos.get('medio_cielo', {}).get('texto', 'N/A')}")
            
            # Casas
            doc.add_heading('🏠 Cúspides de Casas (Placidus)', level=1)
            
            casas_table = doc.add_table(rows=1, cols=2)
            casas_table.style = 'Light Grid Accent 1'
            hdr_cells = casas_table.rows[0].cells
            hdr_cells[0].text = 'Casa'
            hdr_cells[1].text = 'Cúspide'
            
            for casa in self.casas:
                row_cells = casas_table.add_row().cells
                row_cells[0].text = f"Casa {casa['numero']}"
                row_cells[1].text = casa['texto']
            
            doc.add_page_break()
            
            # Análisis
            doc.add_heading('📖 Análisis Psico-Astrológico', level=1)
            for parrafo in self.analysis.split('\n\n'):
                if parrafo.strip():
                    doc.add_paragraph(parrafo.strip())
            
            # Footer
            doc.add_paragraph()
            footer = doc.add_paragraph(f"Informe generado el {datetime.now().strftime('%d/%m/%Y %H:%M')}")
            footer.runs[0].font.size = Pt(9)
            footer.runs[0].font.color.rgb = RGBColor(128, 128, 128)
            footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Guardar en buffer
            buffer = BytesIO()
            try:
                doc.save(buffer)
                buffer.seek(0)
                return buffer
            except Exception as save_err:
                print(f"❌ Error guardando DOCX: {save_err}", file=sys.stderr)
                import traceback
                traceback.print_exc(file=sys.stderr)
                raise Exception(f"Error al guardar el documento DOCX: {str(save_err)}")
        except Exception as e:
            print(f"❌ Error en generate_docx: {type(e).__name__}: {e}", file=sys.stderr)
            import traceback
            traceback.print_exc(file=sys.stderr)
            raise


def generate_report(carta_data: Dict, format: str, analysis_text: Optional[str] = None, nombre: str = ""):
    """
    Función principal para generar informes en cualquier formato
    
    Args:
        carta_data: Datos de la carta astral
        format: Formato deseado ('pdf', 'docx', 'markdown', 'html')
        analysis_text: Texto del análisis (opcional)
        nombre: Nombre del consultante (para portada)
        
    Returns:
        Contenido del informe en el formato solicitado
    """
    generator = ReportGenerator(carta_data, analysis_text, nombre)
    
    format_lower = format.lower()
    
    if format_lower == 'markdown' or format_lower == 'md':
        return generator.generate_markdown()
    elif format_lower == 'html' or format_lower == 'web':
        return generator.generate_html()
    elif format_lower == 'pdf':
        return generator.generate_pdf()
    elif format_lower == 'docx' or format_lower == 'doc':
        return generator.generate_docx()
    else:
        raise ValueError(f"Formato no soportado: {format}. Use: pdf, docx, markdown, html")

